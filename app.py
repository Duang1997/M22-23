import os
import streamlit as st
from docx import Document
from docx.shared import Inches
from PIL import Image

# กำหนดค่าหน้าจอ Streamlit
st.set_page_config(
    page_title="ระบบบันทึกภาพและจัดทำแฟ้มผู้ต้องหา", page_icon="📋", layout="centered"
)

st.title("ระบบบันทึกภาพและข้อมูลผู้ต้องหา")
st.markdown("---")

# ส่วนที่ 1: การกรอกข้อมูลส่วนบุคคล
st.subheader("1. ข้อมูลผู้ต้องหา")
col1, col2 = st.columns(2)
with col1:
  first_name = st.text_input("ชื่อ")
  national_id = st.text_input(
      "เลขประจำตัวประชาชน (13 หลัก)", max_chars=13
  ).strip()
with col2:
  last_name = st.text_input("นามสกุล")

st.markdown("---")

# ส่วนที่ 2: การบันทึกภาพ 4 ด้าน (ถ่ายภาพจากกล้อง หรือ อัปโหลดไฟล์)
st.subheader("2. บันทึกภาพถ่าย 4 ด้าน")

angles = {
    "front": "ภาพหน้าตรง",
    "left": "ภาพด้านซ้าย",
    "right": "ภาพด้านขวา",
    "back": "ภาพด้านหลัง",
}

captured_images = {}

for key, label in angles.items():
  st.markdown(f"**{label}**")
  input_method = st.radio(
      f"เลือกวิธีนำเข้าสำหรับ{label}",
      ("ถ่ายภาพจากกล้อง", "อัปโหลดไฟล์ภาพ"),
      key=f"method_{key}",
      horizontal=True,
  )

  if input_method == "ถ่ายภาพจากกล้อง":
    img_file = st.camera_input(f"ถ่าย{label}", key=f"cam_{key}")
  else:
    img_file = st.file_uploader(
        f"เลือกไฟล์ภาพ{label}", type=["jpg", "jpeg", "png"], key=f"file_{key}"
    )

  if img_file is not None:
    captured_images[key] = Image.open(img_file)
  else:
    captured_images[key] = None

  st.markdown("---")


# ฟังก์ชันสร้างเอกสาร Word
def generate_docx(nid, fname, lname, images):
  doc = Document()
  doc.add_heading("บันทึกข้อมูลและภาพถ่ายผู้ต้องหา", 0)

  doc.add_paragraph(f"ชื่อ-นามสกุล: {fname} {lname}")
  doc.add_paragraph(f"เลขประจำตัวประชาชน: {nid}")
  doc.add_paragraph("ภาพถ่าย 4 ด้าน:")

  # แทรกรูปภาพลงในเอกสาร
  label_map = {
      "front": "ภาพหน้าตรง",
      "left": "ภาพด้านซ้าย",
      "right": "ภาพด้านขวา",
      "back": "ภาพด้านหลัง",
  }

  for k, img in images.items():
    if img is not None:
      temp_path = f"temp_{k}.jpg"
      img.save(temp_path)
      doc.add_paragraph(label_map[k])
      doc.add_picture(temp_path, width=Inches(3.0))
      if os.path.exists(temp_path):
        os.remove(temp_path)

  filename = f"ภาพแนบ-{nid}-{fname} {lname}.docx"
  doc.save(filename)
  return filename


# ส่วนที่ 3: ปุ่มบันทึกและดาวน์โหลด
if st.button("บันทึกข้อมูลและสร้างไฟล์รายงาน", type="primary"):
  if not national_id or not first_name or not last_name:
    st.error("กรุณากรอกข้อมูล ชื่อ นามสกุล และเลขประจำตัวประชาชนให้ครบถ้วน")
  elif len(captured_images) == 0:
    st.error("กรุณานำเข้าภาพถ่ายอย่างน้อย 1 ด้าน")
  else:
    with st.spinner("กำลังประมวลผลและสร้างไฟล์เอกสาร..."):
      # สร้างไฟล์ Word
      docx_filename = generate_docx(
          national_id, first_name, last_name, captured_images
      )

      st.success("สร้างเอกสารสำเร็จ")

      # แสดงปุ่มดาวน์โหลดไฟล์ Word
      with open(docx_filename, "rb") as f:
        st.download_button(
            label="ดาวน์โหลดไฟล์ Word (.docx)",
            data=f,
            file_name=docx_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

      # หมายเหตุเรื่อง PDF (กรณีรันบน Cloud แนะนำให้ใช้เครื่องมือแปลงไฟล์เพิ่มเติมหรือดาวน์โหลด Word ไปพิมพ์ตรง)
      st.info(
          "หมายเหตุ: สามารถเปิดไฟล์ Word ที่ดาวน์โหลดเพื่อตรวจสอบหน้ากระดาษและสั่งพิมพ์หรือแปลงเป็น PDF ได้ทันที"
      )